"""Document extractors for Word and HWP files"""

from pathlib import Path
from typing import Optional, Union

from ..utils.korean import normalize_korean
from ..utils.logging import get_logger

logger = get_logger("extractors.doc")


class DocumentExtractor:
    """Extract text from various document formats"""

    def __init__(self):
        self._docx_available = None
        self._olefile_available = None

    def _check_docx(self) -> bool:
        """Check if python-docx is available"""
        if self._docx_available is None:
            try:
                import docx
                self._docx_available = True
            except ImportError:
                self._docx_available = False
                logger.warning("python-docx not installed. DOCX extraction disabled.")
        return self._docx_available

    def _check_olefile(self) -> bool:
        """Check if olefile is available for HWP"""
        if self._olefile_available is None:
            try:
                import olefile
                self._olefile_available = True
            except ImportError:
                self._olefile_available = False
                logger.warning("olefile not installed. HWP extraction disabled.")
        return self._olefile_available

    def extract_docx(self, file_path: Union[str, Path]) -> str:
        """Extract text from DOCX file"""
        if not self._check_docx():
            return ""

        import docx

        file_path = Path(file_path)
        if not file_path.exists():
            logger.error(f"DOCX file not found: {file_path}")
            return ""

        try:
            doc = docx.Document(str(file_path))
            text_parts = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text:
                    text_parts.append(para.text)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text:
                            row_text.append(cell.text)
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            full_text = "\n\n".join(text_parts)
            return normalize_korean(full_text)

        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {e}")
            return ""

    def extract_doc(self, file_path: Union[str, Path]) -> str:
        """Extract text from DOC file (old Word format)"""
        # DOC files are more complex - we'll try olefile for basic extraction
        if not self._check_olefile():
            return ""

        import olefile

        file_path = Path(file_path)
        if not file_path.exists():
            return ""

        try:
            ole = olefile.OleFileIO(str(file_path))

            # Try to find WordDocument stream
            if ole.exists("WordDocument"):
                # Basic text extraction - DOC format is complex
                # For production, consider using antiword or LibreOffice
                data = ole.openstream("WordDocument").read()

                # Try to extract readable text
                # This is a simplified approach
                text = ""
                try:
                    text = data.decode("utf-16-le", errors="ignore")
                except Exception:
                    try:
                        text = data.decode("cp949", errors="ignore")
                    except Exception:
                        text = data.decode("utf-8", errors="ignore")

                # Filter to printable characters
                import string
                printable = set(string.printable + "가-힣")
                text = "".join(c if c in printable or ord(c) > 0xAC00 else " " for c in text)

                ole.close()
                return normalize_korean(text)

            ole.close()
            return ""

        except Exception as e:
            logger.error(f"Error extracting text from DOC {file_path}: {e}")
            return ""

    def extract_hwp(self, file_path: Union[str, Path]) -> str:
        """Extract text from HWP file (Korean word processor)"""
        if not self._check_olefile():
            return ""

        import olefile
        import zlib

        file_path = Path(file_path)
        if not file_path.exists():
            return ""

        try:
            ole = olefile.OleFileIO(str(file_path))

            # HWP files store text in PrvText or BodyText streams
            text_parts = []

            # Try PrvText first (preview text, easier to extract)
            if ole.exists("PrvText"):
                data = ole.openstream("PrvText").read()
                try:
                    text = data.decode("utf-16-le", errors="ignore")
                    text_parts.append(text)
                except Exception:
                    pass

            # Try BodyText sections
            for entry in ole.listdir():
                if entry[0] == "BodyText":
                    try:
                        data = ole.openstream(entry).read()

                        # Try to decompress if needed
                        try:
                            data = zlib.decompress(data, -15)
                        except Exception:
                            pass

                        # HWP uses UTF-16LE
                        text = data.decode("utf-16-le", errors="ignore")

                        # Filter control characters but keep Korean
                        import re
                        text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text)

                        if text.strip():
                            text_parts.append(text)

                    except Exception as e:
                        logger.debug(f"Could not extract {entry}: {e}")

            ole.close()

            full_text = "\n\n".join(text_parts)
            return normalize_korean(full_text)

        except Exception as e:
            logger.error(f"Error extracting text from HWP {file_path}: {e}")
            return ""

    def extract(self, file_path: Union[str, Path]) -> dict:
        """Extract text from any supported document format"""
        file_path = Path(file_path)

        if not file_path.exists():
            return {"text": "", "error": "File not found"}

        ext = file_path.suffix.lower()

        if ext == ".docx":
            text = self.extract_docx(file_path)
        elif ext == ".doc":
            text = self.extract_doc(file_path)
        elif ext in (".hwp", ".hwpx"):
            text = self.extract_hwp(file_path)
        else:
            return {"text": "", "error": f"Unsupported format: {ext}"}

        return {
            "text": text,
            "file_path": str(file_path),
            "file_type": ext[1:],  # Remove dot
        }

    def get_supported_extensions(self) -> list[str]:
        """Get list of supported file extensions"""
        extensions = []

        if self._check_docx():
            extensions.append(".docx")

        if self._check_olefile():
            extensions.extend([".doc", ".hwp", ".hwpx"])

        return extensions
